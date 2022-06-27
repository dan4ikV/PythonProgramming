import re
import numpy as np
import matplotlib.pyplot as plt
import docx
import io
import requests
import shutil
from PIL import Image


def download_img(url, filename):
    r = requests.get(url, stream=True)
    r.raw.decode_content = True

    with open(f"{filename}.jpg", 'wb') as f:
        shutil.copyfileobj(r.raw, f)


def image_magic(final_image_name, url='https://cdnb.artstation.com/p/assets/images/images/019/454/943/large/serhii-livariuk-1-artst.jpg?1563553991'):
    download_img(url, 'background')
    image = Image.open('background.jpg')
    image = image.resize((800, 1000))
    box = (20, 20, 780, 980)
    image = image.crop(box)

    logo = Image.open('logo.png')
    logo = logo.resize((150, 142))
    logo = logo.rotate(180)
    image.paste(logo, (600, 10), mask=logo)
    image.save(f"{final_image_name}.jpg")


def find_same_number_of_words(arr):
    d = {}
    for entry in arr:
        if entry[1] in d:
            d[entry[1]] += (entry[0] + 1,)
        else:
            d[entry[1]] = (entry[0] + 1,)
    return d


def average_number_of_words(arr):
    avg = 0
    for entry in arr:
        avg += entry[1]
    return avg / len(arr)


def build_graph(arr):
    memfile = io.BytesIO()
    fig, ax = plt.subplots()
    x = np.array(list(zip(*arr))[0]) + 1
    y = np.array(list(zip(*arr))[1])
    fig.set_figheight(5)
    fig.set_figwidth(10)
    plt.xlabel('paragraph number')
    plt.ylabel('words')
    plt.xticks(x)
    ax.bar(x, y)
    plt.savefig(memfile)
    return memfile


def preprocess(arr):
    arr = list(zip(range(0, len(arr)), arr))
    arr.sort(key=lambda x: x[1])
    return arr


def create_word_document(paragraph_data, title, image_url, book_author, report_author='Danylo Vasylyshyn 256711'):
    document = docx.Document()

    #styles
    style = document.styles['Normal']
    style.font.name = "Times New Roman"
    style.font.size = docx.shared.Pt(11)
    style.font.bold = False

    #add title heading
    document.add_heading(title, 0)

    #add picture #1
    response = requests.get(image_url, stream=True)
    image = io.BytesIO(response.content)
    pic1 = document.add_picture(image, width=docx.shared.Inches(4.8))
    last_paragraph = document.paragraphs[-1]
    last_paragraph.alignment = 1

    paragraph = document.add_paragraph()
    desc = paragraph.add_run('Taras killed his son Andriy')
    desc.font.name = "Times New Roman"
    desc.font.size = docx.shared.Pt(9)
    desc.font.bold = True
    desc.font.italic = True
    desc.font.color.rgb = docx.shared.RGBColor(100, 100, 100)
    paragraph.alignment = 1


    #adding book and report author
    document.add_heading(f'Book author: {book_author}\nReport author: {report_author}')
    last_paragraph = document.paragraphs[-1]
    last_paragraph.alignment = 1

    #new page
    document.add_page_break()

    #plot
    paragraph = document.add_paragraph()
    header = paragraph.add_run('Plot of distribution of lengths of paragraphs')
    header.font.name = "Times New Roman"
    header.font.size = docx.shared.Pt(18)
    header.font.bold = True
    paragraph.alignment = 1

    pic2 = build_graph(paragraph_data)
    document.add_picture(pic2, width=docx.shared.Inches(6))
    last_paragraph = document.paragraphs[-1]
    last_paragraph.alignment = 1

    #paragraph data
    paragraph2 = document.add_paragraph()
    header = paragraph2.add_run("Chapter 1 data about paragraphs:\n\n")
    header.font.name = "Times New Roman"
    header.font.size = docx.shared.Pt(14)
    header.font.bold = True
    paragraph.alignment = 1

    ret = find_same_number_of_words(paragraph_data)
    for key, value in ret.items():
        if len(value) > 1:
            data = paragraph2.add_run(f'Paragraphs {value} all have the length of {key} words\n')
            data.font.name = "Times New Roman"
            data.font.size = docx.shared.Pt(11)
            data.font.bold = False

    a = list(ret.items())
    data = paragraph2.add_run(f'\nAvaerage number of words in paragraphs of chapter 1: {round(average_number_of_words(paragraph_data), 1)}')
    paragraph2.add_run(f'\nMaximum number of words is {a[-1][0]} in paragraph(s) {a[-1][1]}')
    paragraph2.add_run(f'\nMinimum number of words is {a[0][0]} in paragraph(s) {a[0][1]}')

    document.save('demo.docx')


def do_all():
    author = ""
    title = ""
    paragraph_data = []
    pat2 = re.compile(r"\b([A-Z\s]{3,})\b")

    def get(to_get, line):
        pat = re.compile(f"{to_get}: (.*)")
        if pat.match(line) is not None:
            temp = pat.match(line).groups()
            print(temp[0])
            if to_get == "Author":
                nonlocal author
                author = temp[0]
            elif to_get == "Title":
                nonlocal title
                title = temp[0]

    def words_in_first_paragraph(line):
        if line == "CHAPTER I\n":
            nonlocal read
            read = True
            paragraph_data.append(0)

        if read:
            last_elem = len(paragraph_data) - 1

            if line.isspace(): # if it is the end of the paragraph
                temp = file.__next__()
                if pat2.match(temp) is not None:    #check if second chapter has not yet started
                    nonlocal finished_reading
                    paragraph_data[last_elem] -= paragraph_data[last_elem] % 10;
                    finished_reading = True
                elif not temp.startswith("“") and not temp.isspace():   #check if a paragraph starts with quotation, count it as a continuation of the current one
                    paragraph_data[last_elem] -= paragraph_data[last_elem] % 10;
                    paragraph_data.append(len(temp.split(" ")))
                else:
                    paragraph_data[last_elem] += len(temp.split(" "))

            else:
                paragraph_data[last_elem] += len(line.split(" "))

    #data for reading the file
    read = False
    finished_reading = False
    file = open("1197-0.txt", encoding="UTF-8")

    for line in file:
        get("Author", line)
        get("Title", line)
        if finished_reading is False:
            words_in_first_paragraph(line)

    paragraph_data = preprocess(paragraph_data)
    build_graph(paragraph_data)
    ret = find_same_number_of_words(paragraph_data)

    for key, value in ret.items():
        if len(value) > 1:
            print(f'Paragraphs {value} all have the length of {key} words')

    url = 'https://cdnb.artstation.com/p/assets/images/images/019/454/943/large/serhii-livariuk-1-artst.jpg?1563553991'
    image_magic("composed_image")
    create_word_document(paragraph_data, title, url, author)


do_all()
